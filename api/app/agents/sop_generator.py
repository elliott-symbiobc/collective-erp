"""
sop_generator.py — Stage 5: Claude-powered CRISPR SOP generation.

Takes a genome edit package (list of dicts from edit_prioritizer) and
uses the Claude API to generate a structured Standard Operating Procedure,
then writes it to a DOCX file via python-docx.

Public API
----------
generate_sop(strain_name, substrate_name, edits, output_path) -> str
    Synchronous entry point. Returns path to the written DOCX file.
"""

import json
import logging
import os
from datetime import date
from pathlib import Path
from typing import Any, Optional

import anthropic
from app.agents.citation_config import CITATION_SYSTEM_INSTRUCTION
from app.agents.usage_logger import log_anthropic_call

logger = logging.getLogger(__name__)

SOP_OUTPUT_DIR = Path("/opt/symbio/sops")

_SOP_SYSTEM_CORE = """\
You are a molecular biology protocol writer specialising in CRISPR-Cas9 editing of \
filamentous fungi, specifically Aspergillus species. \
You write clear, step-by-step standard operating procedures that can be executed by a \
trained lab scientist with molecular biology experience. \
You follow ALCOA+ documentation principles: Attributable, Legible, Contemporaneous, \
Original, Accurate, plus Complete, Consistent, Enduring, Available.
"""

SOP_SYSTEM_PROMPT = CITATION_SYSTEM_INSTRUCTION + "\n\n" + _SOP_SYSTEM_CORE

SOP_USER_TEMPLATE = """\
Generate a complete CRISPR-Cas9 Standard Operating Procedure for the following genome \
editing campaign.

STRAIN: {strain_name}
SUBSTRATE: {substrate_name}
DATE: {today}

EDIT TARGETS:
{edit_summary}

Return a JSON object with this exact structure (do not include markdown code fences):
{{
  "title": "SOP title string",
  "sop_number": "GEN-EDIT-XXX",
  "version": "1.0",
  "scope": "One or two sentences describing what this SOP covers.",
  "safety": ["bullet 1", "bullet 2", ...],
  "materials": [
    {{"item": "item name", "specification": "grade/supplier/cat no", "quantity": "amount"}}
  ],
  "sections": [
    {{
      "heading": "Section heading",
      "steps": ["Step 1 text.", "Step 2 text.", ...]
    }}
  ],
  "expected_outcomes": ["outcome 1", "outcome 2"],
  "troubleshooting": [
    {{"problem": "problem description", "cause": "likely cause", "solution": "corrective action"}}
  ],
  "references": ["reference 1", "reference 2"]
}}

Required sections (in this order):
1. Reagent and Equipment Preparation
2. sgRNA Synthesis and Quality Control
3. Cas9 Protein Preparation
4. Protoplast Preparation from {strain_name}
5. RNP Complex Assembly
6. Transformation
7. Selection and Colony Screening
8. PCR Verification of Edit
9. Off-Target Assessment
10. Documentation and Archiving

For each edit target, include the specific sgRNA sequence and HDR template information \
in the relevant sections. Be specific about Aspergillus-specific considerations such as \
protoplast preparation from fungal mycelia and selection on minimal media.
"""


def _build_edit_summary(edits: list[dict]) -> str:
    lines = []
    for i, ed in enumerate(edits, 1):
        sg = ed.get("sgrna_sequence") or "pending design"
        lines.append(
            f"{i}. {ed['edit_type'].upper()} of {ed['target_gene']} "
            f"(feature: {ed['feature_name']}, "
            f"ΔTiter estimate: +{ed.get('delta_titer_estimate', '?')} U/mL)\n"
            f"   sgRNA: {sg}\n"
            f"   Rationale: {ed.get('description', '')}"
        )
    return "\n".join(lines)


def _call_claude(strain_name: str, substrate_name: str, edits: list[dict]) -> dict:
    """Call the Anthropic API and return the parsed SOP JSON."""
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

    edit_summary = _build_edit_summary(edits)
    user_msg = SOP_USER_TEMPLATE.format(
        strain_name=strain_name,
        substrate_name=substrate_name,
        today=str(date.today()),
        edit_summary=edit_summary,
    )

    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=16000,
        system=SOP_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_msg}],
    )
    log_anthropic_call(
        operation="sop",
        model="claude-opus-4-6",
        input_tokens=message.usage.input_tokens,
        output_tokens=message.usage.output_tokens,
    )

    raw = message.content[0].text.strip()
    # Strip possible markdown fences
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1]
    if raw.endswith("```"):
        raw = raw.rsplit("```", 1)[0].strip()

    # Attempt to parse; if truncated, try to recover a partial object
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        # Find the last complete top-level key and truncate there
        # Append minimal closing to make valid JSON
        for suffix in [']}', ']}]', ']}]}', ']}]}]}']:
            try:
                return json.loads(raw + suffix)
            except json.JSONDecodeError:
                pass
        # Last resort: return minimal fallback
        logger.warning("Could not parse Claude JSON response, using fallback structure")
        return {
            "title": f"CRISPR SOP — {strain_name} × {substrate_name}",
            "sop_number": "GEN-EDIT-AUTO",
            "version": "1.0",
            "scope": "Auto-generated CRISPR editing protocol (JSON parsing failed — see raw response).",
            "safety": ["Standard molecular biology PPE required."],
            "materials": [],
            "sections": [{"heading": "Raw Claude Response", "steps": [raw[:2000]]}],
            "expected_outcomes": [],
            "troubleshooting": [],
            "references": [],
        }


def _write_docx(sop: dict, output_path: Path) -> None:
    """Write the SOP JSON structure to a Word document with SBC branding."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    SBC_RED   = RGBColor(0xC3, 0x10, 0x10)  # Thunderbird Red
    DARK_RED  = RGBColor(0x8B, 0x14, 0x14)  # table headers
    SATIN     = RGBColor(0xEE, 0xEC, 0xE1)  # alternating rows
    WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
    DARK_GREY = RGBColor(0x11, 0x18, 0x27)
    MID_GREY  = RGBColor(0x6B, 0x72, 0x80)
    FONT      = 'Calibri'

    def _set_cell_bg(cell, rgb: RGBColor):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), f'{rgb.red:02X}{rgb.green:02X}{rgb.blue:02X}')
        tcPr.append(shd)

    def _add_bottom_border(para, color='C31010', size=6):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(size))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _style_run(run_obj, color: RGBColor, bold=False, size_pt=11):
        run_obj.font.name = FONT
        run_obj.font.color.rgb = color
        run_obj.font.bold = bold
        run_obj.font.size = Pt(size_pt)

    def _add_h1(doc, text, page_break=False):
        p = doc.add_paragraph()
        if page_break:
            run_obj = p.add_run()
            run_obj.add_break(docx_module.enum.text.WD_BREAK.PAGE)
        r = p.add_run(text)
        _style_run(r, RGBColor(0, 0, 0), bold=True, size_pt=20)
        p.paragraph_format.space_before = Pt(28)
        p.paragraph_format.space_after  = Pt(12)
        _add_bottom_border(p)
        return p

    def _add_h2(doc, text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        _style_run(r, SBC_RED, bold=True, size_pt=14)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(8)
        return p

    def _add_body(doc, text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        _style_run(r, DARK_GREY, size_pt=11)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def _add_bullet(doc, text):
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(text)
        _style_run(r, DARK_GREY, size_pt=11)
        return p

    def _style_table(tbl, headers):
        """Apply SBC header styling to a table's first row."""
        hdr_row = tbl.rows[0]
        for ci, cell in enumerate(hdr_row.cells):
            cell.text = ''
            _set_cell_bg(cell, DARK_RED)
            p = cell.paragraphs[0]
            r = p.add_run(str(headers[ci]))
            _style_run(r, WHITE, bold=True, size_pt=10)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
        for ri, row in enumerate(tbl.rows[1:], 0):
            shade = ri % 2 == 1
            for cell in row.cells:
                if shade:
                    _set_cell_bg(cell, SATIN)
                for p in cell.paragraphs:
                    for r in p.runs:
                        _style_run(r, DARK_GREY, size_pt=10)

    # ── Page header / footer ──────────────────────────────────────────────────
    import docx as docx_module
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

    # Header
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hr = hp.add_run('SYMBIO BIOCULINARY  |  SOP')
    _style_run(hr, SBC_RED, bold=True, size_pt=9)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_bottom_border(hp, color='C31010', size=4)

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fr = fp.add_run(f'CONFIDENTIAL — Prepared by Collective ERP, {date.today().year}    Page ')
    _style_run(fr, MID_GREY, size_pt=8)
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run_el = OxmlElement('w:r')
    run_el.append(fld)
    run_el.append(instr)
    run_el.append(fld_end)
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp._p.append(run_el)

    # ── Cover / title ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    brand_p = doc.add_paragraph()
    br = brand_p.add_run('SYMBIO BIOCULINARY')
    _style_run(br, SBC_RED, bold=True, size_pt=14)
    brand_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_p = doc.add_paragraph()
    tr = title_p.add_run(sop.get("title", "CRISPR Editing SOP"))
    _style_run(tr, RGBColor(0, 0, 0), bold=True, size_pt=28)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after  = Pt(14)
    _add_bottom_border(title_p)

    # Metadata table (SOP#, version, date)
    meta = doc.add_table(rows=1, cols=3)
    meta.style = 'Table Grid'
    _style_table(meta, ['SOP Number', 'Version', 'Date'])
    data_row = meta.add_row()
    data_row.cells[0].text = sop.get('sop_number', '')
    data_row.cells[1].text = sop.get('version', '1.0')
    data_row.cells[2].text = str(date.today())
    for cell in data_row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                _style_run(r, DARK_GREY, size_pt=10)
    doc.add_paragraph()

    # ── Scope ─────────────────────────────────────────────────────────────────
    _add_h1(doc, 'Scope')
    _add_body(doc, sop.get('scope', ''))

    # ── Safety ────────────────────────────────────────────────────────────────
    safety = sop.get('safety', [])
    if safety:
        _add_h1(doc, 'Safety and PPE Requirements')
        for item in safety:
            _add_bullet(doc, item)

    # ── Materials ─────────────────────────────────────────────────────────────
    materials = sop.get('materials', [])
    if materials:
        _add_h1(doc, 'Materials and Equipment')
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        _style_table(tbl, ['Item', 'Specification', 'Quantity'])
        for m in materials:
            row = tbl.add_row()
            row.cells[0].text = m.get('item', '')
            row.cells[1].text = m.get('specification', '')
            row.cells[2].text = m.get('quantity', '')
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _style_run(r, DARK_GREY, size_pt=10)
        doc.add_paragraph()

    # ── Procedure sections ────────────────────────────────────────────────────
    for si, section_data in enumerate(sop.get('sections', [])):
        _add_h1(doc, section_data.get('heading', ''), page_break=(si > 0))
        for j, step in enumerate(section_data.get('steps', []), 1):
            p = doc.add_paragraph(style='List Number')
            r = p.add_run(step)
            _style_run(r, DARK_GREY, size_pt=11)

    # ── Expected outcomes ─────────────────────────────────────────────────────
    outcomes = sop.get('expected_outcomes', [])
    if outcomes:
        _add_h1(doc, 'Expected Outcomes')
        for item in outcomes:
            _add_bullet(doc, item)

    # ── Troubleshooting ───────────────────────────────────────────────────────
    ts = sop.get('troubleshooting', [])
    if ts:
        _add_h1(doc, 'Troubleshooting')
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        _style_table(tbl, ['Problem', 'Cause', 'Solution'])
        for t in ts:
            row = tbl.add_row()
            row.cells[0].text = t.get('problem', '')
            row.cells[1].text = t.get('cause', '')
            row.cells[2].text = t.get('solution', '')
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _style_run(r, DARK_GREY, size_pt=10)
        doc.add_paragraph()

    # ── References ────────────────────────────────────────────────────────────
    refs = sop.get('references', [])
    if refs:
        _add_h1(doc, 'References')
        for ref in refs:
            p = doc.add_paragraph(style='List Number')
            r = p.add_run(ref)
            _style_run(r, DARK_GREY, size_pt=11)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def generate_sop(
    strain_name: str,
    substrate_name: str,
    edits: list[dict],
    output_path: Optional[Path] = None,
) -> str:
    """
    Generate a CRISPR SOP DOCX for a list of genome edits.

    Returns the absolute path to the written DOCX file.
    """
    if output_path is None:
        safe_strain = "".join(c if c.isalnum() else "_" for c in strain_name)[:30]
        safe_sub = "".join(c if c.isalnum() else "_" for c in substrate_name)[:20]
        filename = f"SOP_{safe_strain}_{safe_sub}_{date.today()}.docx"
        output_path = SOP_OUTPUT_DIR / filename

    logger.info("Generating SOP for %s × %s (%d edits)", strain_name, substrate_name, len(edits))

    try:
        sop_json = _call_claude(strain_name, substrate_name, edits)
        _write_docx(sop_json, output_path)
        logger.info("SOP written to %s", output_path)
    except Exception as exc:
        logger.error("SOP generation failed: %s", exc)
        raise

    return str(output_path)
